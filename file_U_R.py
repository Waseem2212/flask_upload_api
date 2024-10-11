from flask import Flask,request,jsonify
import os 
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import document,Document
import docx

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER): os.makedirs(UPLOAD_FOLDER)

ALLOWED_EXTENSIONS_PDF = {'pdf'}
ALLOWED_EXTENSIONS_DOC = {'docx'}

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.',1)[1].lower() in allowed_extensions

@app.route('/upload-file', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error':'file not exists'}),400
    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'file not select'}),400
    
    if allowed_file(file.filename, ALLOWED_EXTENSIONS_PDF):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'],filename)
        file.save(file_path)

        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return jsonify({'message':'PDF processed', 'text':text})
    
    elif allowed_file(file.filename, ALLOWED_EXTENSIONS_DOC):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'],filename)
        file.save(file_path)
        text = ''
        # def extract_text(doc):
        #     full_text = []
        #     for para in doc.paragraphs:
        #         full_text.append(para.text)
        #         for table in doc.tables:
        #             for row in table.rows:
        #                 for cell in row.cells:
        #                     full_text.append(cell.text)
        #     return '\n'.join(full_text)
        # text = '\n'.join([para.text for para in Document.paragraphs[0].text])
        # if file_path == 'docx':
        #             text = ''.join([para.text for para in Document(file.filename).paragraphs])

        def getText(filename):
            # print(len(doc.paragraphs))
            doc = docx.Document(filename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)
        text = getText(file_path)
        return jsonify({'message':'word file processed','text':text})
    
    return jsonify({'error':'unsupported file type'}),400

if __name__ == '__main__':
    app.run(debug=True, port=8000)


